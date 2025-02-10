import os
from datetime import datetime
import sqlite3
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def export_to_word_simple(dataset, analysis, output_path):
    """Exporta el análisis de un dataset a un archivo Word con formato mejorado."""

    doc = Document()

    # **Portada con Logo**
    logo_path = os.path.join(os.getcwd(), "Generador_plantillas", "static", "images", "logo2.png")
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar párrafo antes de la imagen
    if os.path.exists(logo_path):
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(2))
    else:
        logo_paragraph.add_run("[LOGO NO ENCONTRADO]")

    title = doc.add_paragraph("\nInforme de Análisis de Datos")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n")

    # **Detalles del Dataset**
    doc.add_heading(f"Análisis del Dataset: {dataset['file_name']}", level=1)
    doc.add_paragraph(f"Categoría: {dataset['category']}")
    doc.add_paragraph(f"Descripción: {dataset['description']}")
    doc.add_paragraph(f"Fecha de Creación: {dataset['created_at']}")
    doc.add_paragraph("\n")

    # **Tabla de Métricas**
    doc.add_heading("Métricas", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'Valor'

    for key, value in dataset['metricas'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(round(value, 2))
    doc.add_paragraph("\n")

    # **Recuperar el análisis de la base de datos**
    with sqlite3.connect('./data/processed_datasets.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT analysis FROM datasets WHERE id = ?", (dataset["id"],))
        analysis_row = cursor.fetchone()

    analysis = analysis_row[0] if analysis_row and analysis_row[0] else "No se encontró el análisis."

    # **Gráficos Asociados**
    doc.add_heading("Gráficos Asociados", level=2)
    for chart_path in dataset['chart_paths'].split(','):
        full_chart_path = os.path.join(os.getcwd(), "Generador_plantillas", "static", "processed_charts", chart_path.strip())
        
        if os.path.exists(full_chart_path):  # **Evita agregar gráficos inexistentes**
            doc.add_paragraph(f"Gráfico: {chart_path.split('/')[-1]}")
            doc.add_picture(full_chart_path, width=Inches(5))
        else:
            doc.add_paragraph("[Imagen no encontrada]")

    doc.add_paragraph("\n")

    # **Análisis Generado**
    doc.add_heading("Análisis", level=2)
    doc.add_paragraph(analysis)

    # **Pie de Página con Fecha**
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.text = f"Generado el: {dataset['created_at']}"
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # **Guardar el Documento**
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)